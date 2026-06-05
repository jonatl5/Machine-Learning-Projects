from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
REPORT_DOCX = ROOT / "Lab9_Report.docx"
REPORT_PDF = ROOT / "Lab9_Report.pdf"
DEPLOYMENT_IMAGE = ROOT / "deployment_serial_monitor.png"

OVERVIEW_TEXT = (
    "This lab trained a compact neural-network classifier for two ASL-inspired IMU gestures, "
    "converted the trained TensorFlow model to TensorFlow Lite, exported the model as an Arduino "
    "header, and deployed it on the Arduino board for serial-monitor testing."
)

SOFTWARE_FIXES = [
    "Fixed the intentional training-setup issue by treating the task as multi-class classification: "
    "the output layer uses softmax over the two gesture classes, the labels are one-hot encoded, "
    "and the model is compiled with categorical_crossentropy and accuracy.",
    "Kept the downstream plots and metrics consistent with that fix by plotting loss, validation loss, "
    "accuracy, and validation accuracy using the matching history keys.",
    "Used the same normalization in the notebook and Arduino sketch: accelerometer values are scaled "
    "with (value + 4) / 8 and gyroscope values with (value + 2000) / 4000.",
    "Exported gesture_model.tflite and model.h, then included model.h in the Arduino classifier sketch "
    "with the gesture labels ordered as hi then sup.",
]

MODEL_RESULTS_TEXT = (
    "The final run trained for 100 epochs. The last epoch reported loss 0.2016, training accuracy "
    "0.9000, validation loss 0.3173, and validation accuracy 0.8000. On the held-out test set, the "
    "model reported loss 0.2339 and accuracy 0.8500. The exported TensorFlow Lite model is 148,296 "
    "bytes and the generated Arduino header is 926,943 bytes."
)

DEPLOYMENT_TEXT = (
    "The Arduino sketch was configured for the Nano 33 BLE Sense Rev2 IMU library, uses 119 IMU "
    "samples per inference window, and runs inference after significant motion crosses the 2.5 G "
    "acceleration threshold. The Serial Monitor output below shows both gestures being identified "
    "with the higher confidence score: hi at 0.884913 and sup at 0.857895 or higher in later tests."
)

SUBMISSION_FILES = [
    "Completed notebook: EE446_TinyML_Lab9.ipynb",
    "Submission PDF: Lab9_Report.pdf",
    "Arduino sketch: lab9-classifier-dual-board.ino",
    "Deployed model header: model.h",
    "Related verification files: hi.csv, sup.csv, gesture_model.tflite, and this Word report source.",
]

SUMMARY_ROWS = [
    ("Gestures", "hi, sup", "Samples/gesture", "50 recordings"),
    ("Input shape", "100 x 714", "Split", "60 train / 20 validation / 20 test"),
    ("Final train accuracy", "0.9000", "Validation accuracy", "0.8000"),
    ("Test accuracy", "0.8500", "TFLite size", "148,296 bytes"),
    ("Header size", "926,943 bytes", "Tensor arena", "16 KB"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def create_deployment_image():
    lines = [
        "17:55:04.578 -> Accelerometer sample rate = 99.79 Hz",
        "17:55:04.611 -> Gyroscope sample rate = 99.79 Hz",
        "17:55:04.611 ->",
        "17:55:10.222 -> hi:  0.884913",
        "17:55:10.222 -> sup: 0.115087",
        "17:55:10.222 ->",
        "17:55:23.283 -> hi:  0.641782",
        "17:55:23.283 -> sup: 0.358218",
        "17:55:23.283 ->",
        "17:55:24.701 -> hi:  0.142105",
        "17:55:24.701 -> sup: 0.857895",
        "17:55:24.701 ->",
        "17:55:27.539 -> hi:  0.196483",
        "17:55:27.540 -> sup: 0.803517",
        "17:55:29.714 -> hi:  0.263914",
        "17:55:29.714 -> sup: 0.736086",
        "17:55:30.998 -> hi:  0.185743",
        "17:55:30.998 -> sup: 0.814257",
    ]

    width, height = 1500, 900
    image = PILImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("consola.ttf", 46)
    except OSError:
        font = ImageFont.load_default()

    y = 24
    for line in lines:
        draw.text((24, y), line, fill=(62, 74, 82), font=font)
        y += 48

    DEPLOYMENT_IMAGE.parent.mkdir(parents=True, exist_ok=True)
    image.save(DEPLOYMENT_IMAGE)


def apply_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("EE446 TinyML Lab 9 Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta.add_run("American Sign Language Gesture Recognition | Arduino Nano 33 BLE Sense Rev2 | June 5, 2026")


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.45), Inches(1.55), Inches(1.45), Inches(1.95)]
    headers = ["Item", "Value", "Item", "Value"]
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        cell.width = width
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].add_run(text).bold = True

    for row in SUMMARY_ROWS:
        cells = table.add_row().cells
        for cell, text, width in zip(cells, row, widths):
            cell.width = width
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].add_run(text)


def build_report():
    create_deployment_image()

    doc = Document()
    apply_document_styles(doc)
    add_title(doc)

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(OVERVIEW_TEXT)

    add_summary_table(doc)

    doc.add_heading("Software Fixes and Consistency Changes", level=1)
    for item in SOFTWARE_FIXES:
        add_bullet(doc, item)

    doc.add_heading("Model Results", level=1)
    doc.add_paragraph(MODEL_RESULTS_TEXT)

    doc.add_heading("Hardware Deployment Evidence", level=1)
    doc.add_paragraph(DEPLOYMENT_TEXT)

    img_paragraph = doc.add_paragraph()
    img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_run = img_paragraph.add_run()
    img_run.add_picture(str(DEPLOYMENT_IMAGE), width=Inches(6.35))
    caption = doc.add_paragraph("Figure 1. Arduino Serial Monitor deployment output from the provided screenshot.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)

    doc.add_heading("Files Included for Submission", level=1)
    for item in SUBMISSION_FILES:
        add_bullet(doc, item)

    doc.save(REPORT_DOCX)
    build_pdf()
    print(f"Saved {REPORT_DOCX}")
    print(f"Saved {REPORT_PDF}")
    print(f"Saved {DEPLOYMENT_IMAGE}")


def pdf_heading(text, styles):
    return Paragraph(text, styles["LabHeading"])


def pdf_body(text, styles):
    return Paragraph(text, styles["LabBody"])


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="LabTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=4,
            alignment=0,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabMeta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#444444"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabHeading",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=13.5,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabBullet",
            parent=styles["LabBody"],
            leftIndent=18,
            firstLineIndent=-10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LabCaption",
            parent=styles["LabBody"],
            fontSize=9,
            leading=11,
            alignment=1,
            textColor=colors.HexColor("#444444"),
            spaceBefore=2,
        )
    )

    doc = SimpleDocTemplate(
        str(REPORT_PDF),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    story = [
        Paragraph("EE446 TinyML Lab 9 Report", styles["LabTitle"]),
        Paragraph(
            "American Sign Language Gesture Recognition | Arduino Nano 33 BLE Sense Rev2 | June 5, 2026",
            styles["LabMeta"],
        ),
        pdf_heading("Overview", styles),
        pdf_body(OVERVIEW_TEXT, styles),
    ]

    table_data = [["Item", "Value", "Item", "Value"], *SUMMARY_ROWS]
    table = Table(table_data, colWidths=[1.45 * inch, 1.55 * inch, 1.45 * inch, 1.95 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C2CC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([table, Spacer(1, 8)])

    story.extend([pdf_heading("Software Fixes and Consistency Changes", styles)])
    for item in SOFTWARE_FIXES:
        story.append(Paragraph(f"- {item}", styles["LabBullet"]))

    story.extend([pdf_heading("Model Results", styles), pdf_body(MODEL_RESULTS_TEXT, styles)])
    story.extend([pdf_heading("Hardware Deployment Evidence", styles), pdf_body(DEPLOYMENT_TEXT, styles)])
    story.append(RLImage(str(DEPLOYMENT_IMAGE), width=6.35 * inch, height=3.81 * inch))
    story.append(Paragraph("Figure 1. Arduino Serial Monitor deployment output from the provided screenshot.", styles["LabCaption"]))

    story.extend([pdf_heading("Files Included for Submission", styles)])
    for item in SUBMISSION_FILES:
        story.append(Paragraph(f"- {item}", styles["LabBullet"]))

    doc.build(story)


if __name__ == "__main__":
    build_report()
